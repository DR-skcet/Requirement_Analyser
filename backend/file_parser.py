import os
import pytesseract
from docx import Document
from openpyxl import load_workbook
from PIL import Image
from PyPDF2 import PdfReader
from typing import Tuple
import pdfplumber
import mimetypes
import tempfile
import csv

from fastapi import FastAPI, UploadFile, File

import requests
import json

app = FastAPI()

# Load Groq API key from env
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
GROQ_API_URL = "https://api.groq.com/openai/v1/chat/completions"


def extract_text_from_file(filename: str, content: bytes) -> Tuple[str, dict]:
    mime_type, _ = mimetypes.guess_type(filename)
    print(f"[INFO] Detected MIME type: {mime_type} for file: {filename}")

    with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(filename)[1]) as tmp:
        tmp.write(content)
        tmp_path = tmp.name

    text = ""
    metadata = {}

    try:
        if filename.endswith(".pdf") or mime_type == 'application/pdf':
            print("[INFO] Parsing PDF file")
            text = extract_text_from_pdf(tmp_path)
            metadata = extract_pdf_metadata(tmp_path)

        elif filename.endswith(".docx") or mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            print("[INFO] Parsing DOCX file")
            text = extract_text_from_docx(tmp_path)
            metadata = extract_docx_metadata(tmp_path)

        elif filename.endswith(".xlsx") or mime_type == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet':
            print("[INFO] Parsing XLSX file")
            text = extract_text_from_xlsx(tmp_path)

        elif filename.endswith(".csv") or mime_type == 'text/csv':
            print("[INFO] Parsing CSV file")
            text = extract_text_from_csv(tmp_path)

        elif mime_type and mime_type.startswith('image/'):
            print("[INFO] Parsing image file")
            text = extract_text_from_image(tmp_path)

        elif mime_type == 'text/plain' or filename.endswith(".txt"):
            print("[INFO] Parsing TXT file")
            text = extract_text_from_txt(tmp_path)

        else:
            print("[WARNING] Unsupported file format or unknown MIME type")
            text = "[ERROR] Unsupported file format"

    except Exception as e:
        print(f"[ERROR] Failed to parse file: {e}")
        text = "[ERROR] Could not extract content"

    finally:
        os.unlink(tmp_path)

    return text, metadata


def extract_text_from_pdf(file_path):
    with pdfplumber.open(file_path) as pdf:
        return "\n".join(page.extract_text() or "" for page in pdf.pages)


def extract_pdf_metadata(file_path: str) -> dict:
    reader = PdfReader(file_path)
    doc_info = reader.metadata
    metadata = {
        "author": doc_info.author if doc_info.author else None,
        "creator": doc_info.creator if doc_info.creator else None,
        "producer": doc_info.producer if doc_info.producer else None,
        "subject": doc_info.subject if doc_info.subject else None,
        "title": doc_info.title if doc_info.title else None,
        "creation_date": doc_info.creation_date if doc_info.creation_date else None,
        "modification_date": doc_info.modification_date if doc_info.modification_date else None,
    }
    return metadata


def extract_text_from_docx(file_path):
    doc = Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])


def extract_docx_metadata(file_path: str) -> dict:
    doc = Document(file_path)
    core_properties = doc.core_properties
    metadata = {
        "author": core_properties.author,
        "created": core_properties.created,
        "last_modified_by": core_properties.last_modified_by,
        "modified": core_properties.modified,
        "title": core_properties.title,
        "subject": core_properties.subject,
    }
    return metadata


def extract_text_from_xlsx(file_path):
    wb = load_workbook(file_path)
    content = ""
    for sheet in wb:
        for row in sheet.iter_rows(values_only=True):
            content += " ".join([str(cell) if cell else "" for cell in row]) + "\n"
    return content


def extract_text_from_csv(file_path):
    content = ""
    with open(file_path, newline='', encoding='utf-8') as csvfile:
        reader = csv.reader(csvfile)
        for row in reader:
            content += " ".join(row) + "\n"
    return content


def extract_text_from_image(file_path):
    img = Image.open(file_path)
    return pytesseract.image_to_string(img)


def extract_text_from_txt(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()


def extract_requirements(text: str) -> dict:
    if not text.strip():
        return {"error": "Empty input text"}

    headers = {
        "Authorization": f"Bearer {GROQ_API_KEY}",
        "Content-Type": "application/json"
    }

    prompt = f"""
You are a requirements extraction assistant.

Extract the following from the given software requirements text:

- Functional Requirements
- Non-Functional Requirements
- Constraints
- Assumptions

Return the output as a JSON object with keys exactly as above. Return only the JSON, nothing else.

Text:
\"\"\"
{text}
\"\"\"
"""

    payload = {
        "model": "llama3-70b-8192",
        "messages": [
            {
                "role": "system",
                "content": "You are a helpful assistant specialized in software requirements extraction."
            },
            {
                "role": "user",
                "content": prompt
            }
        ],
        "temperature": 0,
        "max_tokens": 2048
    }

    try:
        response = requests.post(GROQ_API_URL, json=payload, headers=headers)
        response.raise_for_status()
        data = response.json()

        content = data["choices"][0]["message"]["content"].strip()

        try:
            result = json.loads(content)
        except json.JSONDecodeError:
            return {"error": "Failed to parse LLM output as JSON", "raw_output": content}

        return result

    except requests.exceptions.RequestException as e:
        return {"error": str(e)}
    except (KeyError, IndexError):
        return {"error": "Unexpected response format", "raw": data}


# Example chunk_text function - split text into chunks of ~1000 words or less
def chunk_text(text: str, max_words=1000):
    words = text.split()
    chunks = []
    for i in range(0, len(words), max_words):
        chunks.append(" ".join(words[i:i+max_words]))
    return chunks


@app.post("/upload/")
async def upload_file(file: UploadFile = File(...)):
    content = await file.read()
    extracted_text, metadata = extract_text_from_file(file.filename, content)
    text_chunks = chunk_text(extracted_text)

    all_requirements = {
        "Functional Requirements": [],
        "Non-Functional Requirements": [],
        "Constraints": [],
        "Assumptions": []
    }

    for chunk in text_chunks:
        result = extract_requirements(chunk)
        if "error" not in result:
            for key in all_requirements.keys():
                # Expect each value to be a list, if not, convert to list
                value = result.get(key, [])
                if isinstance(value, list):
                    all_requirements[key].extend(value)
                elif isinstance(value, str) and value.strip():
                    all_requirements[key].append(value.strip())

    return {
        "metadata": metadata,
        "requirements": all_requirements,
    }
